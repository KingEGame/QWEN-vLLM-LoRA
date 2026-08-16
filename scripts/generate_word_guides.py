#!/usr/bin/env python3
"""Generate the repository's concise operator Word guides."""

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "word"

NAVY = "17365D"
BLUE = "2F75B5"
PALE_BLUE = "D9EAF7"
LIGHT = "F2F5F8"
WHITE = "FFFFFF"
TEXT = "202A35"
MUTED = "5B6573"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=55, start=70, bottom=55, end=70) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Page ")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])
    run.add_text(" of ")
    begin2 = OxmlElement("w:fldChar")
    begin2.set(qn("w:fldCharType"), "begin")
    instr2 = OxmlElement("w:instrText")
    instr2.set(qn("xml:space"), "preserve")
    instr2.text = "NUMPAGES"
    end2 = OxmlElement("w:fldChar")
    end2.set(qn("w:fldCharType"), "end")
    run._r.extend([begin2, instr2, end2])


def new_doc(short_title: str) -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.48)
    section.bottom_margin = Inches(0.48)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.2)
    section.footer_distance = Inches(0.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(9)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.space_after = Pt(2.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE

    for name, size, color, before, after in (
        ("Title", 19, NAVY, 0, 4),
        ("Subtitle", 9, MUTED, 0, 5),
        ("Heading 1", 12, NAVY, 5, 2),
        ("Heading 2", 9.8, BLUE, 3, 1.5),
    ):
        style = styles[name]
        style.font.name = "Aptos Display" if name in {"Title", "Heading 1"} else "Aptos"
        style.font.size = Pt(size)
        style.font.bold = name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.text = f"QWEN-vLLM-LoRA  |  {short_title}"
    header.style = styles["Caption"]
    header.runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    header.runs[0].font.size = Pt(7.5)
    add_page_number(section.footer.paragraphs[0])
    section.footer.paragraphs[0].runs[0].font.size = Pt(7.5)
    section.footer.paragraphs[0].runs[0].font.color.rgb = RGBColor.from_string(MUTED)
    return doc


def title(doc: Document, text: str, subtitle: str) -> None:
    doc.add_paragraph(text, style="Title")
    doc.add_paragraph(subtitle, style="Subtitle")


def p(doc: Document, text: str = "", *, bold_prefix: str | None = None, style=None):
    para = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        para.add_run(bold_prefix).bold = True
        para.add_run(text[len(bold_prefix):])
    else:
        para.add_run(text)
    return para


def bullet(doc: Document, text: str, level: int = 0) -> None:
    para = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    para.paragraph_format.left_indent = Inches(0.18 + level * 0.16)
    para.paragraph_format.first_line_indent = Inches(-0.12)
    para.paragraph_format.space_after = Pt(1)
    para.add_run(text)


def code(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, LIGHT)
    set_cell_margins(cell, 55, 90, 55, 90)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(text.splitlines()):
        if i:
            para.add_run().add_break()
        run = para.add_run(line)
        run.font.name = "Cascadia Mono"
        run.font.size = Pt(7.6)
        run.font.color.rgb = RGBColor.from_string(TEXT)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths=None) -> None:
    t = doc.add_table(rows=1, cols=len(headers))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for idx, h in enumerate(headers):
        cell = t.rows[0].cells[idx]
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor.from_string(WHITE)
        run.font.size = Pt(7.8)
    for row_idx, values in enumerate(rows):
        cells = t.add_row().cells
        for idx, value in enumerate(values):
            shade(cells[idx], WHITE if row_idx % 2 == 0 else LIGHT)
            set_cell_margins(cells[idx])
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            cells[idx].paragraphs[0].paragraph_format.space_after = Pt(0)
            run = cells[idx].paragraphs[0].add_run(value)
            run.font.size = Pt(7.65)
    if widths:
        for row in t.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def callout(doc: Document, heading: str, body: str) -> None:
    t = doc.add_table(rows=1, cols=1)
    cell = t.cell(0, 0)
    shade(cell, PALE_BLUE)
    set_cell_margins(cell, 70, 95, 70, 95)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    r = para.add_run(heading + "  ")
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(NAVY)
    para.add_run(body)


def page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(0)
    para.add_run().add_break(WD_BREAK.PAGE)


def save(doc: Document, filename: str) -> None:
    props = doc.core_properties
    props.author = "QWEN-vLLM-LoRA repository"
    props.subject = "Repository operations and model training guide"
    props.keywords = "Qwen, vLLM, AWQ, LoRA, QLoRA"
    doc.save(OUT / filename)


def guide_server() -> None:
    doc = new_doc("Server & API")
    title(doc, "Run, Access, and Call the Model Server", "Two-page runbook • Linux/WSL2 + NVIDIA GPU • Repository state: 10 August 2026")
    callout(doc, "Outcome", "Start the Qwen3.6-27B AWQ model locally, verify it, call the OpenAI-compatible API, and select a LoRA adapter when needed.")
    doc.add_heading("1. One-time setup", level=1)
    p(doc, "Run from the repository root. Native Windows does not host this GPU stack; the Windows wrappers enter WSL2. The setup recreates .venv, installs requirements, and verifies CUDA/vLLM.")
    code(doc, "# Windows PowerShell or cmd\nscripts\\setup.cmd\n\n# Linux / already in WSL2\n./scripts/setup.sh\nsource .venv/bin/activate")
    p(doc, "WSL only: if Triton/FlashInfer reports compiler or CUDA-header problems, run bash scripts/_install_usergcc.sh once. Model weights download on first use; keep ample disk space and Hugging Face access.")
    doc.add_heading("2. Start and verify the base server", level=1)
    code(doc, "source .venv/bin/activate\n./scripts/start_server.sh\n# Leave this terminal open. Wait until the server reports it is ready.")
    p(doc, "The script reads config/model.env and currently serves shawnw3i/Qwen3.6-27B-AWQ-MTP on port 8000 with AWQ, 4,096-token context, max 32 concurrent sequences, and 92% GPU-memory utilization.")
    code(doc, "# In a second WSL/Linux terminal\nsource .venv/bin/activate\ncurl http://localhost:8000/v1/models\npython scripts/test_client.py --prompt \"Explain what this server does.\"")
    doc.add_heading("3. How clients reach it", level=1)
    table(doc, ["Client", "Base URL", "Notes"], [
        ["Same WSL/Linux host", "http://localhost:8000/v1", "Preferred test path."],
        ["Windows host with WSL2", "http://localhost:8000/v1", "WSL localhost forwarding normally makes this work."],
        ["Another LAN machine", "http://<server-ip>:8000/v1", "Only after binding/routing and firewall rules are intentionally configured."],
    ], [1.3, 2.1, 3.9])
    callout(doc, "Security", "This configuration has no meaningful API key enforcement. Do not expose port 8000 to the public Internet. Put authentication and TLS in a reverse proxy before any shared deployment.")

    page_break(doc)
    title(doc, "Requests, Adapters, and Recovery", "Use the model name exactly as advertised by /v1/models or by the LoRA launch command.")
    doc.add_heading("4. Call the chat-completions endpoint", level=1)
    code(doc, "curl http://localhost:8000/v1/chat/completions \\\n+  -H \"Content-Type: application/json\" \\\n+  -d '{\"model\":\"shawnw3i/Qwen3.6-27B-AWQ-MTP\",\"messages\":[{\"role\":\"user\",\"content\":\"Give one GPU troubleshooting tip.\"}],\"temperature\":0.2}'")
    code(doc, "from openai import OpenAI\nclient = OpenAI(base_url=\"http://localhost:8000/v1\", api_key=\"not-needed\")\nr = client.chat.completions.create(\n    model=\"shawnw3i/Qwen3.6-27B-AWQ-MTP\",\n    messages=[{\"role\": \"user\", \"content\": \"Hello\"}],\n)\nprint(r.choices[0].message.content)")
    doc.add_heading("5. Serve and call adapters", level=1)
    code(doc, "# Default support adapter: output/lora_adapter\n./scripts/serve_with_lora.sh\npython scripts/test_client.py --model support-adapter\n\n# Two named personal adapters\nLORA_MODULES=\"question-sharper=output/lora_question_sharper,me-assistant=output/lora_me_assistant\" ./scripts/serve_with_lora.sh\npython scripts/test_client.py --model me-assistant --prompt \"Why did vLLM OOM?\"")
    p(doc, "Adapter names become API model names. For the two-stage workflow, run python scripts/personal_pipeline.py \"<messy technical question>\"; it calls question-sharper first and me-assistant second.")
    doc.add_heading("6. Stop, tune, and diagnose", level=1)
    table(doc, ["Symptom", "Action"], [
        ["Connection refused", "Server is not ready, wrong port, or terminal exited. Check its log and /v1/models."],
        ["CUDA OOM / Mamba cache error", "Stop other GPU jobs; reduce MAX_MODEL_LEN and then MAX_NUM_SEQS in config/model.env. Restart."],
        ["Adapter not found", "Confirm its directory contains adapter_config.json and adapter_model.safetensors; verify LORA_MODULES path/name."],
        ["Training needed", "Stop the server with Ctrl+C and confirm nvidia-smi shows VRAM released before training."],
        ["Need extra vLLM flags", "Launch with EXTRA_ARGS=\"--enforce-eager\" ./scripts/start_server.sh (example)."],
    ], [2.0, 5.3])
    callout(doc, "Operational check", "A successful HTTP connection is not enough: verify a sensible response, log latency/errors, and keep a small fixed prompt set for base-versus-adapter regression tests.")
    save(doc, "01-server-run-access-api.docx")


def guide_assets_training() -> None:
    doc = new_doc("Assets & Training")
    title(doc, "Repository Assets, Training, and New Adapters", "Two-page inventory and workflow • Paths are relative to the repository root unless shown otherwise")
    callout(doc, "Current state", "The checkout contains the AWQ-serving configuration, dense-base QLoRA trainer, generated/reviewed datasets, and three trained LoRA adapters.")
    doc.add_heading("1. What is here and where", level=1)
    table(doc, ["Area", "Important locations", "Purpose"], [
        ["Configuration", "config/model.env\nconfig/personal_sources.env", "Serving model/port/memory; dense TRAIN_MODEL; personal-source paths."],
        ["Serve & test", "scripts/start_server.sh\nscripts/serve_with_lora.sh\nscripts/test_client.py", "Base vLLM; vLLM plus named adapters; one request smoke test."],
        ["Train", "scripts/train_lora.py", "Validates JSONL, loads dense 27B as NF4, trains rank-16 PEFT LoRA."],
        ["FAQ data", "data/source_docs/\ndata/generated/raw_qa.jsonl\ndata/train.jsonl", "Source text; model-generated draft; human-approved training set."],
        ["Personal data", "data/personal/candidates/\nquestion_sharp.jsonl\nme_assistant.jsonl", "Extracted candidates; reviewed sharper and assistant datasets."],
        ["Adapters", "output/lora_adapter/\noutput/lora_question_sharper/\noutput/lora_me_assistant/", "Each has adapter config, tokenizer files, and ~159 MB adapter_model.safetensors."],
        ["Guidance/tests", "docs/guides/\ntests/", "Architecture/operator explanations and CPU-only pipeline validation."],
    ], [1.25, 2.6, 3.45])
    p(doc, "Base weights are cached outside the repo in WSL: ~/.cache/huggingface/hub/models--shawnw3i--Qwen3.6-27B-AWQ-MTP/ (serving, about 19 GB) and models--Qwen--Qwen3.6-27B/ (dense training source, about 54+ GB; the audited cache was about 65 GB with partial files). output/ and private JSONL are gitignored.")
    doc.add_heading("2. Dataset contract", level=1)
    code(doc, "{\"instruction\": \"User request\", \"response\": \"Ideal assistant answer\"}")
    bullet(doc, "One valid UTF-8 JSON object per line; both fields must be non-empty strings.")
    bullet(doc, "Review generated or extracted examples. Remove secrets, duplicates, weak answers, copied-input ‘rewrites,’ and contradictory preferences.")
    bullet(doc, "Keep a held-out evaluation set; do not train on every example you use to judge quality.")
    code(doc, "python scripts/validate_dataset.py path/to/new_dataset.jsonl")

    page_break(doc)
    title(doc, "Train and Deploy a New Adapter", "The server and trainer compete for the same GPU: never run both during a 27B training job.")
    doc.add_heading("3. Reuse the existing FAQ pipeline", level=1)
    code(doc, "# With base AWQ server running\npython scripts/generate_training_data.py\n# Review data/generated/raw_qa.jsonl, then promote approved rows\ncp data/generated/raw_qa.jsonl data/train.jsonl\npython scripts/validate_dataset.py data/train.jsonl\n# Stop vLLM (Ctrl+C), confirm VRAM is free\nnvidia-smi\nMAX_SEQ_LENGTH=1024 BATCH_SIZE=1 NUM_EPOCHS=1 GRADIENT_ACCUMULATION_STEPS=8 GRADIENT_CHECKPOINTING=1 python scripts/train_lora.py")
    doc.add_heading("4. Create any new adapter", level=1)
    code(doc, "python scripts/validate_dataset.py data/my_task.jsonl\nMAX_SEQ_LENGTH=1024 BATCH_SIZE=1 NUM_EPOCHS=1 \\\n+  GRADIENT_ACCUMULATION_STEPS=8 GRADIENT_CHECKPOINTING=1 \\\n+  python scripts/train_lora.py --data data/my_task.jsonl --output output/lora_my_task\nLORA_MODULES=\"my-task=output/lora_my_task\" ./scripts/serve_with_lora.sh\npython scripts/test_client.py --model my-task --prompt \"A held-out test request\"")
    p(doc, "Start with a one-epoch smoke run. Compare held-out prompts against the base model; only then increase epochs or data. Keep adapter_config.json with adapter_model.safetensors and record the base model revision, dataset version, parameters, loss, and evaluation result.")
    doc.add_heading("5. Resource budget and safe starting point", level=1)
    table(doc, ["Resource", "Serving AWQ", "27B QLoRA training"], [
        ["GPU VRAM", "Designed for one 24 GB-class NVIDIA GPU; ~19 GB weights plus KV cache/runtime.", "24 GB proven for smoke runs at seq 1024, batch 1, accumulation 8; checkpointing adds safety."],
        ["System RAM", "32 GB practical; more helps loading/cache behavior.", "32 GB WSL + swap is tight but demonstrated; 64 GB host RAM is safer."],
        ["Disk", "Allow ~25 GB for AWQ plus environment/cache headroom.", "Allow ~70 GB dense cache plus environment, datasets, checkpoints; ~100 GB free is a sensible floor."],
        ["Time", "Startup includes model load; first start includes download.", "Strongly dataset/GPU dependent; smoke first, measure, then schedule full runs."],
    ], [1.2, 3.05, 3.05])
    callout(doc, "Important config detail", "train_lora.py currently reads MAX_SEQ_LENGTH, BATCH_SIZE, NUM_EPOCHS, GRADIENT_ACCUMULATION_STEPS, and GRADIENT_CHECKPOINTING. The TRAIN_MAX_SEQ_LENGTH / TRAIN_BATCH_SIZE / TRAIN_NUM_EPOCHS / TRAIN_GRAD_ACCUM keys in model.env are not wired to those knobs; use the names shown above.")
    save(doc, "02-assets-training-new-adapters.docx")


def guide_rationale() -> None:
    doc = new_doc("Design Rationale")
    title(doc, "Why AWQ Serving and QLoRA Training?", "Two-page engineering rationale • Current target: Qwen3.6-27B on one 24 GB-class NVIDIA GPU")
    callout(doc, "Decision in one sentence", "Use an inference-optimized 4-bit AWQ checkpoint to serve, but load the matching dense checkpoint as training-oriented 4-bit NF4 to learn small LoRA adapters.")
    doc.add_heading("1. AWQ instead of dense BF16 for serving", level=1)
    p(doc, "A 27B model in BF16 needs roughly 54 GB for parameters alone (27 billion × 2 bytes), before KV cache and runtime workspaces. It cannot fit on the target 24 GB GPU. The AWQ checkpoint is about 19 GB on disk and keeps weights at 4-bit for inference, leaving limited but usable VRAM for context, concurrency, and LoRA execution.")
    table(doc, ["Choice", "Benefit", "Cost / when to choose otherwise"], [
        ["AWQ 4-bit", "Fits the capable 27B model; fast vLLM inference; dynamic named LoRAs.", "Small quantization-quality loss; tight KV cache. Use BF16 only with ~60+ GB VRAM or multi-GPU capacity."],
        ["Dense BF16", "Highest numerical fidelity and simplest merge/training source.", "Far beyond this GPU for 27B serving; higher hardware and power cost."],
        ["Smaller dense model", "Fast experiments, larger context/batch on the same GPU.", "Lower base capability, but often the best dataset-development baseline."],
    ], [1.25, 3.0, 3.05])
    doc.add_heading("2. QLoRA instead of ordinary BF16 LoRA", level=1)
    p(doc, "LoRA defines the trainable low-rank matrices; QLoRA is the memory-saving way the frozen base is loaded during training. This trainer quantizes the matching dense Qwen/Qwen3.6-27B base to NF4 with double quantization and BF16 compute. Only the LoRA parameters receive updates. The output is a normal PEFT LoRA adapter—not a special QLoRA-only format.")
    bullet(doc, "Ordinary BF16 LoRA still keeps the frozen 27B base in BF16, so the base alone exceeds 24 GB VRAM.")
    bullet(doc, "AWQ is optimized for inference, not this bitsandbytes/PEFT training path; train from the dense source instead.")
    bullet(doc, "At inference, vLLM applies the floating-point LoRA delta over the quantized AWQ base without permanently merging it.")
    callout(doc, "Compatibility condition", "The dense training source and AWQ serving checkpoint must represent the same underlying base: architecture, layer names/shapes, tokenizer behavior, and pretrained revision. Structural loading success is necessary; held-out quality testing is still required.")

    page_break(doc)
    title(doc, "Why These Values—and What to Improve", "Tune one variable at a time and judge with held-out examples, latency, throughput, and peak VRAM.")
    doc.add_heading("3. Current parameter choices", level=1)
    table(doc, ["Value", "Why it is reasonable", "If larger / smaller"], [
        ["AWQ group size 128", "Checkpoint’s established accuracy/size compromise.", "Not a runtime knob; re-quantization is required to change it."],
        ["Context 4,096", "Leaves KV-cache room on 24 GB while covering common support/personal prompts.", "8,192 improves long tasks but raises KV memory; smaller improves stability/throughput."],
        ["Max sequences 32", "Throughput-oriented ceiling for short requests.", "Lower for OOM/latency consistency; raise only after measured load tests."],
        ["GPU utilization 0.92", "Lets vLLM reserve most VRAM for weights/cache.", "Higher risks fragmentation/OOM; lower is safer but reduces cache capacity."],
        ["LoRA rank/alpha 16/16", "~159 MB adapters; enough capacity for style/task behavior without large overhead.", "Higher rank may learn complex shifts but costs memory/overfits; lower is cheaper but less expressive."],
        ["Targets: q/k/v/o + gate/up/down", "Adapts attention and MLP paths across the model.", "Fewer targets reduce size/compute but can limit behavior change."],
        ["Dropout 0; LR 2e-4", "Common efficient SFT starting point for a small adapter.", "Dropout may help small/noisy data; lower LR is safer if loss/behavior is unstable."],
        ["Seq 1,024; batch 1; accum 8", "Fits proven 24 GB smoke runs; effective batch = 8 examples/update.", "Longer/bigger raises activation VRAM; smaller context truncates examples."],
        ["1 epoch smoke; 3 configured", "One proves the pipeline; more passes may fit style/data.", "More can memorize tiny datasets; choose by validation, not training loss alone."],
    ], [1.7, 3.15, 2.45])
    doc.add_heading("4. Highest-value improvements", level=1)
    bullet(doc, "Data first: manually fix copied-input question-sharper rows, deduplicate, redact secrets, and add held-out task/style evaluations. Better examples usually beat larger rank or more epochs.")
    bullet(doc, "Reproducibility: pin dense and AWQ model revisions; save dataset hashes, seed, package versions, command, GPU, loss, and evaluation scores with every adapter.")
    bullet(doc, "Configuration: wire the TRAIN_* keys in model.env to the trainer or remove them; currently they can imply settings that are not used.")
    bullet(doc, "Reliability/security: benchmark context × concurrency, add readiness/metrics, and require authenticated TLS ingress before remote access.")
    bullet(doc, "Iteration speed: develop data and evaluation on a 4B–8B model, then confirm the final adapter on 27B. Enable gradient checkpointing when memory is marginal.")
    callout(doc, "When to change architecture", "Use cloud/multi-GPU BF16 LoRA or full fine-tuning only when evaluation proves QLoRA capacity/quantization is the bottleneck. Use retrieval for changing facts; use prompting when behavior does not need persistent learned adaptation.")
    save(doc, "03-awq-qlora-parameter-rationale.docx")


def main() -> None:
    OUT.mkdir(parents=True, exist_ok=True)
    guide_server()
    guide_assets_training()
    guide_rationale()
    print(f"Generated 3 guides in {OUT}")


if __name__ == "__main__":
    main()
